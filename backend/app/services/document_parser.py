"""Парсинг договоров и заказ-нарядов в структурированные таблицы. Поддержаны
Excel (xlsx/xlsm/xls), CSV, OpenDocument (ods), Word (docx) и PDF — у разных
контрагентов и СТО документооборот ведётся в разных форматах, и требовать
от пользователя конвертировать всё в один формат вручную неудобно.

Каждая строка приводится к общему формату:
    {"article": str | None, "name": str, "qty": float | None, "price": float | None}

Эвристики распознавания колонок написаны широко (RU-синонимы), потому что
формат договоров у разных контрагентов не стандартизован. Позиции, которые
не удалось уверенно распарсить, всё равно возвращаются — с article=None —
чтобы matcher.py мог попытаться сопоставить их через LLM, а ReviewMatches
показал их человеку.
"""

from __future__ import annotations

import logging
import os
import re
from functools import lru_cache

import pandas as pd

logger = logging.getLogger(__name__)

# Списки не претендуют на исчерпывающую полноту всех формулировок, что
# бывают у поставщиков (заказчик не ограничится уже присланными файлами) —
# сравнение по подстроке (см. _match_column/_find_export_column) и так уже
# ловит большинство вариаций ("Цена с НДС, руб." совпадёт с алиасом "цена"
# без отдельной записи) — здесь только те КОРНИ, которых иначе не было бы
# ни в одном алиасе вовсе.
ARTICLE_COLUMN_ALIASES = ["артикул", "article", "код", "sku", "парт", "part_number", "характеристик"]
NAME_COLUMN_ALIASES = ["наименование", "название", "name", "описание", "товар", "запчасть", "номенклатур", "позиция", "изделие", "предмет закупки"]
QTY_COLUMN_ALIASES = ["кол-во", "количество", "qty", "quantity", "шт"]
PRICE_COLUMN_ALIASES = ["цена", "price", "стоимость", "сумма", "тариф"]

# Таблица ставок за нормо-час по маркам ТС (см. parse_hourly_rate_table
# ниже) — колонки ищутся по алиасам заголовка, как и везде в этом файле.
# Реальный файл заказчика (приложение к тендерному контракту, .docx) кладёт
# марку и модель ОДНОЙ колонкой "Марка (модель)" — алиас "марка" ловит её
# по подстроке без отдельной записи. MODEL_COLUMN_ALIASES — на случай, если
# в файле марка и модель всё же разнесены по разным колонкам.
MAKE_COLUMN_ALIASES = ["марка", "make", "brand", "марка тс", "марка автомобиля"]
MODEL_COLUMN_ALIASES = ["модель"]
RATE_COLUMN_ALIASES = ["ставка", "цена н/ч", "цена нормо-часа", "стоимость н/ч", "цена", "rate", "price", "стоимость"]
# Строка "ИТОГО ..."/"Всего ..." в конце такой таблицы — это сумма по
# таблице, а не реальная ставка (реальный файл заказчика: "ИТОГО с учетом
# аукционного снижения (55%): ... 6209.99" — своя "марка" и число, похожее
# на ставку, но это агрегат, который выдал бы мусорную ставку в 6000+ ₽).
_TOTAL_ROW_MARKERS = ("итого", "всего")

# Колонки печатной формы заказ-наряда 1С ("Выполненные работы"/"Расходная
# накладная", см. parse_repair_order_export ниже) — по алиасам заголовка, а
# не по фиксированной позиции. Раньше колонки читались строго по индексу
# (row[9]/row[10]/row[12] и т.п.) — это совпадало с ОДНИМ конкретным
# шаблоном отчёта 1С (реальный файл заказчика), но другая конфигурация 1С
# или версия отчёта вполне может расставить колонки иначе — тогда позиционное
# чтение молча подставляло бы норму часов в цену или наоборот, вместо того
# чтобы честно не найти колонку.
# Переиспользуем те же синонимы, что и generic-парсер выше (ARTICLE/NAME/
# QTY/PRICE_COLUMN_ALIASES), а не заводим более узкий отдельный список —
# "Наименование" в этой печатной форме означает то же самое, что и в любом
# другом договоре/накладной. Добавляем только то, чего там нет: "№ кат." —
# аббревиатура именно этого отчёта 1С для артикула/кода.
LABOR_CATALOG_CODE_ALIASES = ARTICLE_COLUMN_ALIASES + ["№ кат"]
LABOR_DESCRIPTION_ALIASES = NAME_COLUMN_ALIASES + ["работа", "операция"]
LABOR_HOURLY_RATE_ALIASES = ["цена н/ч", "цена нормо-часа", "цена за час", "стоимость н/ч"]
LABOR_NORM_HOURS_ALIASES = ["норма", "нормо-час", "нормочас"]
LABOR_TOTAL_ALIASES = ["всего", "итого", "сумма"]

MATERIAL_ARTICLE_ALIASES = ARTICLE_COLUMN_ALIASES + ["№ кат"]
MATERIAL_NAME_ALIASES = NAME_COLUMN_ALIASES
MATERIAL_QTY_ALIASES = QTY_COLUMN_ALIASES
MATERIAL_PRICE_ALIASES = PRICE_COLUMN_ALIASES


class DocumentParseError(RuntimeError):
    pass


def _match_column(columns: list[str], aliases: list[str]) -> str | None:
    normalized = {c: str(c).strip().lower() for c in columns if not str(c).strip().lower().startswith("unnamed")}
    for col, norm in normalized.items():
        if any(alias in norm for alias in aliases):
            return col
    return None


def _dataframe_to_lines(df: pd.DataFrame) -> list[dict]:
    columns = list(df.columns)
    article_col = _match_column(columns, ARTICLE_COLUMN_ALIASES)
    name_col = _match_column(columns, NAME_COLUMN_ALIASES)
    qty_col = _match_column(columns, QTY_COLUMN_ALIASES)
    price_col = _match_column(columns, PRICE_COLUMN_ALIASES)

    if name_col is None:
        raise DocumentParseError(f"Не удалось найти колонку с наименованием среди {columns}")

    lines = []
    for _, row in df.iterrows():
        name = row.get(name_col)
        if pd.isna(name) or not str(name).strip():
            continue
        lines.append(
            {
                "article": _clean(row.get(article_col)) if article_col else None,
                "name": str(name).strip(),
                "qty": _to_float(row.get(qty_col)) if qty_col else None,
                "price": _to_float(row.get(price_col)) if price_col else None,
            }
        )
    return lines


def _split_make_model(text: str) -> tuple[str, str | None]:
    """"Hyundai Accent" -> ("Hyundai", "Accent"); "Chevrolet" -> ("Chevrolet", None).
    Эвристика (первое слово — марка, остаток — модель) верна для подавляющего
    большинства марок ("Toyota Land Cruiser", "Hyundai Santa Fe"); двусловные
    марки без модели в одной ячейке ("Great Wall") — известное исключение,
    которое эта эвристика не отличит от "марка+модель"."""
    parts = text.split(None, 1)
    if len(parts) == 2:
        return parts[0], parts[1].strip()
    return parts[0], None


def _expand_make_cell(make_text: str, value: float, model_text: str | None = None) -> list[dict]:
    """Одна "ячейка" с маркой(-ами) + число (ставка в рублях ИЛИ норма в
    часах — вызывающий код сам знает, что это такое) -> одна или несколько
    строк {vehicle_make, vehicle_model, value}. Если модель уже известна
    отдельно (своя колонка, или её уже выделила LLM при OCR-разборе) —
    доверяем ей как есть. Иначе разбираем make_text сами: одна ячейка может
    перечислять сразу несколько марок/моделей с одинаковым числом через
    запятую (реальный файл заказчика: "Renault Sandero, Nissan Almera
    Classik, ..., Hyundai Accent, Hyundai Sonata" — все по одной цене)."""
    if model_text:
        return [{"vehicle_make": make_text, "vehicle_model": model_text, "value": value}]

    lines = []
    for segment in make_text.split(","):
        segment = segment.strip()
        if not segment:
            continue
        make, model = _split_make_model(segment)
        lines.append({"vehicle_make": make, "vehicle_model": model, "value": value})
    return lines


def _dataframe_to_rate_lines(df: pd.DataFrame) -> list[dict]:
    columns = list(df.columns)
    make_col = _match_column(columns, MAKE_COLUMN_ALIASES)
    model_col = _match_column(columns, MODEL_COLUMN_ALIASES)
    rate_col = _match_column(columns, RATE_COLUMN_ALIASES)
    if model_col == make_col:
        # "Марка (модель)" одной колонкой matches оба алиаса сразу — это НЕ
        # отдельная колонка модели, а комбинированная ячейка (см.
        # _expand_make_cell).
        model_col = None

    if make_col is None:
        raise DocumentParseError(f"Не удалось найти колонку с маркой ТС среди {columns}")
    if rate_col is None:
        raise DocumentParseError(f"Не удалось найти колонку со ставкой за нормо-час среди {columns}")

    lines = []
    for _, row in df.iterrows():
        make_cell = row.get(make_col)
        if pd.isna(make_cell) or not str(make_cell).strip():
            continue
        make_text = str(make_cell).strip()
        if any(marker in make_text.lower() for marker in _TOTAL_ROW_MARKERS):
            continue

        rate = _to_float(row.get(rate_col))
        if rate is None or rate <= 0:
            continue

        model_text = _clean(row.get(model_col)) if model_col else None
        for expanded in _expand_make_cell(make_text, rate, model_text):
            lines.append(
                {
                    "vehicle_make": expanded["vehicle_make"],
                    "vehicle_model": expanded["vehicle_model"],
                    "hourly_rate": expanded["value"],
                }
            )
    return lines


def parse_hourly_rate_table(file_path: str, llm_client=None) -> list[dict]:
    """Таблица ставок за нормо-час по маркам/моделям ТС (для контрагента или
    договора, см. app/services/hourly_rate_import.py). Каждая строка:
    {"vehicle_make": str, "vehicle_model": str | None, "hourly_rate": float}.

    Поддержаны и простые таблицы (xlsx/xls/ods/csv — одна колонка "Марка",
    одна "Ставка"), и печатная форма приложения к тендерному контракту
    (docx — реальный файл заказчика: таблица "Марка (модель) | Цена" внутри
    Word-документа, возможно несколько таблиц в файле, включая пустые
    служебные — из них просто не наберётся ни одной строки, это не ошибка),
    и PDF с текстовым слоем — тем же путём, что и docx.

    Скан/фото (jpg/png) и PDF БЕЗ текстового слоя (сфотографированное
    бумажное приложение к тендеру — вполне реальный случай, не только
    цифровой оригинал) идут через OCR + LLM-извлечение — тот же путь, что
    и распознавание сканов заказ-нарядов/договоров (см. services/ocr.py,
    LLMClient.extract_table_from_text). Для этого пути нужен llm_client —
    если он не передан, а файл оказался сканом, вернётся понятная ошибка,
    а не тихая пустота."""
    from app.services.ocr import is_image_extension

    ext = os.path.splitext(file_path)[1].lower()
    if is_image_extension(ext):
        return _parse_rate_table_via_ocr(file_path, llm_client)

    if ext in (".xlsx", ".xlsm", ".xls"):
        dataframes = [_read_excel_df(file_path)]
    elif ext == ".ods":
        dataframes = [_read_ods_df(file_path)]
    elif ext == ".csv":
        dataframes = [_read_csv_df(file_path)]
    elif ext == ".docx":
        dataframes = extract_docx_tables(file_path)
    elif ext == ".pdf":
        dataframes = extract_pdf_tables(file_path)
    else:
        raise DocumentParseError(f"Неподдерживаемый формат файла для таблицы ставок: {ext}")

    lines: list[dict] = []
    last_error: DocumentParseError | None = None
    for df in dataframes:
        try:
            lines.extend(_dataframe_to_rate_lines(df))
        except DocumentParseError as exc:
            last_error = exc
    if lines:
        return lines

    if ext == ".pdf":
        # Ни одна таблица не нашлась текстом — похоже на скан без
        # текстового слоя, а не настоящий PDF-документ с таблицей.
        return _parse_rate_table_via_ocr(file_path, llm_client)

    raise last_error or DocumentParseError("В файле не найдено ни одной таблицы со ставками")


def _parse_rate_table_via_ocr(file_path: str, llm_client) -> list[dict]:
    if llm_client is None:
        raise DocumentParseError(
            "Файл похож на скан/фото — чтобы распознать таблицу ставок, нужна выбранная LLM-модель"
        )

    from app.services.ocr import OcrError, extract_text

    try:
        raw_text = extract_text(file_path)
    except OcrError as exc:
        raise DocumentParseError(str(exc)) from exc
    if not raw_text.strip():
        raise DocumentParseError("Не удалось распознать текст в файле (пустой результат OCR)")

    try:
        rows = llm_client.extract_table_from_text(raw_text, ["vehicle_make", "vehicle_model", "hourly_rate"])
    except Exception as exc:
        raise DocumentParseError(f"Не удалось извлечь таблицу ставок из распознанного текста: {exc}") from exc

    lines: list[dict] = []
    for row in rows:
        make_text = _clean(row.get("vehicle_make"))
        if not make_text:
            continue
        if any(marker in make_text.lower() for marker in _TOTAL_ROW_MARKERS):
            continue
        rate = _to_float(row.get("hourly_rate"))
        if rate is None or rate <= 0:
            continue
        model_text = _clean(row.get("vehicle_model"))
        for expanded in _expand_make_cell(make_text, rate, model_text):
            lines.append(
                {
                    "vehicle_make": expanded["vehicle_make"],
                    "vehicle_model": expanded["vehicle_model"],
                    "hourly_rate": expanded["value"],
                }
            )

    if not lines:
        raise DocumentParseError("Не удалось найти ни одной ставки в распознанном тексте")
    return lines


def _dataframe_to_labor_catalog_lines(df: pd.DataFrame) -> list[dict]:
    columns = list(df.columns)
    make_col = _match_column(columns, MAKE_COLUMN_ALIASES)
    model_col = _match_column(columns, MODEL_COLUMN_ALIASES)
    operation_col = _match_column(columns, LABOR_DESCRIPTION_ALIASES)
    norm_hours_col = _match_column(columns, LABOR_NORM_HOURS_ALIASES)
    if model_col == make_col:
        model_col = None  # "Марка (модель)" одной колонкой — см. _dataframe_to_rate_lines

    if make_col is None:
        raise DocumentParseError(f"Не удалось найти колонку с маркой ТС среди {columns}")
    if operation_col is None:
        raise DocumentParseError(f"Не удалось найти колонку с операцией среди {columns}")
    if norm_hours_col is None:
        raise DocumentParseError(f"Не удалось найти колонку с нормо-часами среди {columns}")

    lines = []
    for _, row in df.iterrows():
        make_text = _clean(row.get(make_col))
        operation_text = _clean(row.get(operation_col))
        if not make_text or not operation_text:
            continue
        if any(marker in make_text.lower() for marker in _TOTAL_ROW_MARKERS):
            continue

        norm_hours = _to_float(row.get(norm_hours_col))
        if norm_hours is None or norm_hours <= 0:
            continue

        model_text = _clean(row.get(model_col)) if model_col else None
        # Марка одной ячейкой может перечислять несколько марок/моделей через
        # запятую (см. _expand_make_cell) — та же операция и норма относится
        # к каждой из них.
        for expanded in _expand_make_cell(make_text, norm_hours, model_text):
            lines.append(
                {
                    "vehicle_make": expanded["vehicle_make"],
                    "vehicle_model": expanded["vehicle_model"],
                    "operation_name": operation_text,
                    "norm_hours": expanded["value"],
                }
            )
    return lines


def parse_labor_catalog_table(file_path: str, llm_client=None) -> list[dict]:
    """Таблица справочника нормо-часов (операция + норма часов по маркам/
    моделям ТС, см. app/models/labor_catalog.py) — тот же набор форматов и
    те же эвристики распознавания колонок, что и parse_hourly_rate_table
    (см. её докстринг), только вместо ставки в рублях — норма в часах, и
    обязательна колонка с названием операции. Каждая строка:
    {"vehicle_make": str, "vehicle_model": str | None, "operation_name": str,
    "norm_hours": float}."""
    from app.services.ocr import is_image_extension

    ext = os.path.splitext(file_path)[1].lower()
    if is_image_extension(ext):
        return _parse_labor_catalog_via_ocr(file_path, llm_client)

    if ext in (".xlsx", ".xlsm", ".xls"):
        dataframes = [_read_excel_df(file_path)]
    elif ext == ".ods":
        dataframes = [_read_ods_df(file_path)]
    elif ext == ".csv":
        dataframes = [_read_csv_df(file_path)]
    elif ext == ".docx":
        dataframes = extract_docx_tables(file_path)
    elif ext == ".pdf":
        dataframes = extract_pdf_tables(file_path)
    else:
        raise DocumentParseError(f"Неподдерживаемый формат файла для справочника нормо-часов: {ext}")

    lines: list[dict] = []
    last_error: DocumentParseError | None = None
    for df in dataframes:
        try:
            lines.extend(_dataframe_to_labor_catalog_lines(df))
        except DocumentParseError as exc:
            last_error = exc
    if lines:
        return lines

    if ext == ".pdf":
        return _parse_labor_catalog_via_ocr(file_path, llm_client)

    raise last_error or DocumentParseError("В файле не найдено ни одной строки с нормо-часами")


def _parse_labor_catalog_via_ocr(file_path: str, llm_client) -> list[dict]:
    if llm_client is None:
        raise DocumentParseError(
            "Файл похож на скан/фото — чтобы распознать справочник нормо-часов, нужна выбранная LLM-модель"
        )

    from app.services.ocr import OcrError, extract_text

    try:
        raw_text = extract_text(file_path)
    except OcrError as exc:
        raise DocumentParseError(str(exc)) from exc
    if not raw_text.strip():
        raise DocumentParseError("Не удалось распознать текст в файле (пустой результат OCR)")

    try:
        rows = llm_client.extract_table_from_text(
            raw_text, ["vehicle_make", "vehicle_model", "operation_name", "norm_hours"]
        )
    except Exception as exc:
        raise DocumentParseError(f"Не удалось извлечь таблицу нормо-часов из распознанного текста: {exc}") from exc

    lines: list[dict] = []
    for row in rows:
        make_text = _clean(row.get("vehicle_make"))
        operation_text = _clean(row.get("operation_name"))
        if not make_text or not operation_text:
            continue
        if any(marker in make_text.lower() for marker in _TOTAL_ROW_MARKERS):
            continue
        norm_hours = _to_float(row.get("norm_hours"))
        if norm_hours is None or norm_hours <= 0:
            continue
        model_text = _clean(row.get("vehicle_model"))
        for expanded in _expand_make_cell(make_text, norm_hours, model_text):
            lines.append(
                {
                    "vehicle_make": expanded["vehicle_make"],
                    "vehicle_model": expanded["vehicle_model"],
                    "operation_name": operation_text,
                    "norm_hours": expanded["value"],
                }
            )

    if not lines:
        raise DocumentParseError("Не удалось найти ни одной строки с нормо-часами в распознанном тексте")
    return lines


def _clean(value) -> str | None:
    if value is None or pd.isna(value):
        return None
    return str(value).strip() or None


def _to_float(value) -> float | None:
    if value is None or pd.isna(value):
        return None
    try:
        return float(str(value).replace(",", ".").replace(" ", ""))
    except ValueError:
        return None


def _read_excel_df(file_path: str) -> pd.DataFrame:
    # .xls — старый бинарный формат (OLE2), openpyxl умеет только xlsx/xlsm
    # и молча падает на нём — нужен отдельный движок xlrd.
    ext = os.path.splitext(file_path)[1].lower()
    engine = "xlrd" if ext == ".xls" else "openpyxl"
    return pd.read_excel(file_path, engine=engine, dtype=str)


def _read_ods_df(file_path: str) -> pd.DataFrame:
    return pd.read_excel(file_path, engine="odf", dtype=str)


def _read_csv_df(file_path: str) -> pd.DataFrame:
    """Экспорт из 1С/банк-клиентов чаще всего — Windows-1251 и разделитель
    ';', а не запятая — поэтому не полагаемся на дефолты pandas.

    ';' пробуем ЯВНО первым, а не сразу автоопределение (sep=None): у него
    сносит крышу, стоит запятой встретиться где-то в самих данных — например
    в заголовке "Цена, руб." или в ячейке "Renault Sandero, Nissan Almera,
    ..." (реальный случай — марки через запятую в одной ячейке ставок). На
    маленькой выборке sep=None иногда решает, что разделитель — запятая, и
    падает на первой же строке с "лишней" запятой ("Expected N fields...
    saw M"), хотя реальный разделитель во всём файле — ';'."""
    last_error: Exception | None = None
    for encoding in ("utf-8-sig", "cp1251"):
        for sep in (";", None):
            try:
                df = pd.read_csv(file_path, sep=sep, engine="python", encoding=encoding, dtype=str)
            except (UnicodeDecodeError, pd.errors.ParserError) as exc:
                last_error = exc
                continue
            if df.shape[1] > 1:
                return df
            last_error = last_error or ValueError("не удалось определить разделитель")
    raise DocumentParseError(f"Не удалось прочитать CSV (кодировка/разделитель): {last_error}")


def parse_excel(file_path: str) -> list[dict]:
    return _dataframe_to_lines(_read_excel_df(file_path))


def parse_ods(file_path: str) -> list[dict]:
    return _dataframe_to_lines(_read_ods_df(file_path))


def parse_csv(file_path: str) -> list[dict]:
    return _dataframe_to_lines(_read_csv_df(file_path))


def extract_docx_tables(file_path: str) -> list[pd.DataFrame]:
    from docx import Document

    document = Document(file_path)
    tables: list[pd.DataFrame] = []
    for table in document.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        if len(rows) < 2:
            continue
        header, *body = rows
        tables.append(pd.DataFrame(body, columns=header))
    return tables


def extract_pdf_tables(file_path: str) -> list[pd.DataFrame]:
    import pdfplumber

    tables: list[pd.DataFrame] = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            for table in page.extract_tables():
                if not table or len(table) < 2:
                    continue
                header, *rows = table
                tables.append(pd.DataFrame(rows, columns=header))
    return tables


def _tables_to_lines(tables: list[pd.DataFrame]) -> list[dict]:
    """Docx/PDF нередко содержат НЕСКОЛЬКО таблиц на документ (сама
    ведомость запчастей + служебные таблицы — итоги, реквизиты, шапка) —
    раньше единственная таблица без узнаваемой колонки "наименование"
    (см. _dataframe_to_lines) роняла DocumentParseError и обрывала разбор
    ВСЕГО документа, хотя нужная таблица дальше могла распознаться
    нормально. Пропускаем нераспознанные таблицы, а не документ целиком —
    и поднимаем исключение, только если НИ ОДНА таблица не дала строк
    (тогда парсинг закономерно уходит в LLM-фоллбэк, см.
    parse_document_with_ocr_fallback)."""
    all_lines: list[dict] = []
    last_error: DocumentParseError | None = None
    for df in tables:
        try:
            all_lines.extend(_dataframe_to_lines(df))
        except DocumentParseError as exc:
            last_error = exc
            continue
    if not all_lines and last_error is not None:
        raise last_error
    return all_lines


def parse_docx(file_path: str) -> list[dict]:
    tables = extract_docx_tables(file_path)
    if not tables:
        raise DocumentParseError("В документе не найдено ни одной таблицы")
    return _tables_to_lines(tables)


def parse_pdf(file_path: str) -> list[dict]:
    try:
        stat = os.stat(file_path)
    except OSError:
        stat = None
    tables = _parse_pdf_cached(
        file_path,
        stat.st_mtime_ns if stat else 0,
        stat.st_size if stat else 0,
    )
    if not tables:
        raise DocumentParseError("В PDF не найдено ни одной таблицы")
    return [dict(row) for row in tables]


@lru_cache(maxsize=32)
def _parse_pdf_cached(file_path: str, mtime_ns: int, size: int) -> tuple[dict, ...]:
    """Кеширует тяжёлое извлечение PDF в рамках процесса.

    needs_ocr() и parse_document() вызываются последовательно для одного
    файла и раньше дважды проходили все страницы. Ключ по размеру и mtime
    не позволяет использовать старый результат после изменения файла.
    """
    tables = extract_pdf_tables(file_path)
    return tuple(_tables_to_lines(tables))


_PARSERS = {
    ".xlsx": parse_excel,
    ".xlsm": parse_excel,
    ".xls": parse_excel,
    ".ods": parse_ods,
    ".csv": parse_csv,
    ".docx": parse_docx,
    ".pdf": parse_pdf,
}


def parse_document(file_path: str) -> list[dict]:
    ext = os.path.splitext(file_path)[1].lower()
    parser = _PARSERS.get(ext)
    if parser is None:
        raise DocumentParseError(f"Неподдерживаемый формат файла: {ext}")
    return parser(file_path)


def needs_ocr(file_path: str) -> bool:
    from app.services.ocr import is_image_extension

    ext = os.path.splitext(file_path)[1].lower()
    if is_image_extension(ext):
        return True
    if ext == ".pdf":
        try:
            return len(parse_pdf(file_path)) == 0
        except Exception:
            return True
    return False


def _raw_text_for_llm_fallback(file_path: str) -> str:
    """Сырой текст файла ЧИТАЕМОГО формата (не картинка/скан) — для случая,
    когда жёсткий разбор по колонкам (_dataframe_to_lines/_match_column) не
    нашёл ни одной таблицы с узнаваемой шапкой: нестандартные названия
    колонок, шапка не в первой строке, экзотическая вёрстка и т.п.
    Возвращает "" (не бросает исключение), если и это не удалось — тогда
    вызывающий код поднимает исходную ошибку жёсткого парсера, а не эту."""
    ext = os.path.splitext(file_path)[1].lower()
    try:
        if ext in (".xlsx", ".xlsm", ".xls"):
            return _read_excel_df(file_path).to_csv(index=False)
        if ext == ".ods":
            return _read_ods_df(file_path).to_csv(index=False)
        if ext == ".csv":
            for encoding in ("utf-8-sig", "cp1251"):
                try:
                    with open(file_path, encoding=encoding) as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue
            return ""
        if ext == ".docx":
            from docx import Document

            document = Document(file_path)
            parts = [p.text for p in document.paragraphs if p.text.strip()]
            for table in document.tables:
                for row in table.rows:
                    parts.append(" | ".join(cell.text.strip() for cell in row.cells))
            return "\n".join(parts)
        if ext == ".pdf":
            import pdfplumber

            with pdfplumber.open(file_path) as pdf:
                return "\n\n".join(page.extract_text() or "" for page in pdf.pages)
    except Exception:
        logger.warning("Не удалось получить сырой текст %s для LLM-фоллбэка", file_path, exc_info=True)
        return ""
    return ""


def _extract_via_llm(raw_text: str, llm_client, fields: list[str]) -> list[dict]:
    if not raw_text.strip():
        raise DocumentParseError("Не удалось получить текст файла для распознавания")

    try:
        rows = llm_client.extract_table_from_text(raw_text, fields)
    except Exception as exc:
        raise DocumentParseError(f"Не удалось извлечь таблицу из текста: {exc}") from exc

    numeric_fields = {"qty", "price", "stock_qty", "ordered_qty", "reserved_qty", "in_production_qty", "norm_hours"}
    cleaned = []
    for row in rows:
        if not any(row.values()):
            continue
        cleaned.append(
            {
                field: (_to_float(value) if field in numeric_fields else _clean(value))
                for field, value in row.items()
            }
        )
    return cleaned


def parse_document_with_ocr_fallback(file_path: str, llm_client, fields: list[str]) -> list[dict]:
    if not needs_ocr(file_path):
        try:
            return parse_document(file_path)
        except DocumentParseError:
            # Формат читаемый (xlsx/docx/csv/pdf с таблицами), но структура
            # не совпала ни с одним известным вариантом шапки — не сдаёмся
            # сразу, пробуем более гибкое извлечение через LLM по сырому
            # тексту файла, тем же путём, что и для сканов/фото ниже.
            # Раньше LLM-фоллбэк включался только по РАСШИРЕНИЮ файла
            # (картинка/скан-PDF), а не по факту "жёсткий разбор не
            # справился" — нестандартный, но вполне читаемый xlsx/docx
            # заказчика просто падал с "не удалось найти колонку".
            raw_text = _raw_text_for_llm_fallback(file_path)
            if not raw_text.strip():
                raise
            return _extract_via_llm(raw_text, llm_client, fields)

    from app.services.ocr import OcrError, extract_text

    try:
        raw_text = extract_text(file_path)
    except OcrError as exc:
        raise DocumentParseError(str(exc)) from exc
    if not raw_text.strip():
        raise DocumentParseError("Не удалось распознать текст в файле (пустой результат OCR)")
    return _extract_via_llm(raw_text, llm_client, fields)


_ORDER_RE = re.compile(r"Заказ-наряд\s*№\s*(\S+)\s*от\s*(\d{2}\.\d{2}\.\d{4})")
# "гос. номер: ..." между маркой и VIN есть не во всех выгрузках (см.
# testdata/repair_order_1_final.xlsx — там просто "Автомобиль: MAKE MODEL
# VIN: ..."), поэтому этот кусок необязательный.
_VEHICLE_RE = re.compile(r"Автомобиль\s*:\s*(.+?)\s+(?:гос\.\s*номер\s*:\s*(\S+)\s+)?VIN\s*:\s*(\S+)")
# Год либо "год вып. 2011", либо просто "2011 г." (см. тот же файл).
_YEAR_RE = re.compile(r"год\s*вып\.?\s*(\d{4})|(\d{4})\s*г\.?\b")


def _find_export_column(header_row: list, aliases: list[str]) -> int | None:
    """Индекс колонки в СЫРОЙ (без имён pandas) строке заголовка печатной
    формы 1С — по алиасам, см. комментарий у LABOR_*/MATERIAL_*_ALIASES."""
    for idx, cell in enumerate(header_row):
        if not isinstance(cell, str):
            continue
        norm = cell.strip().lower()
        if any(alias in norm for alias in aliases):
            return idx
    return None


def _export_cell(row: list, idx: int | None):
    if idx is None or idx >= len(row):
        return None
    return row[idx]


def _find_number_column(row: list) -> int | None:
    """Позиция колонки "№" в строке заголовка раздела — НЕ всегда row[1]:
    выгрузки 1С обычно вставляют пустую колонку A перед номером (тогда "№"
    в колонке B/index 1), но встречаются и формы без неё, где "№" сразу в
    колонке A/index 0 (см. testdata/repair_order_1_final.xlsx). Ищем по
    всей строке, а не по фиксированному индексу."""
    for idx, cell in enumerate(row):
        if _clean(cell) == "№":
            return idx
    return None


def _looks_like_data_row(row: list, key_col: int | None) -> bool:
    """Отличает первую строку РЕАЛЬНЫХ данных от необязательной строки-
    легенды с номерами колонок ("1 2 3 ... 9"), которую некоторые печатные
    формы 1С вставляют сразу под заголовком раздела, а некоторые — нет (см.
    testdata/repair_order_1_final.xlsx, где данные идут сразу после
    заголовка). В строке-легенде интересующая нас колонка (наименование
    работы/запчасти) содержит такой же короткий номер, как и остальные —
    отличить можно только по тому, что там не текст, а число."""
    if key_col is None:
        return False
    value = _clean(_export_cell(row, key_col))
    if not value:
        return False
    return not value.replace(".", "", 1).isdigit()


def parse_repair_order_export(file_path: str) -> dict | None:
    ext = os.path.splitext(file_path)[1].lower()
    if ext not in (".xlsx", ".xlsm"):
        return None

    df = pd.read_excel(file_path, sheet_name=0, header=None, dtype=str, engine="openpyxl")
    rows = df.values.tolist()

    has_labor_section = any(any(isinstance(c, str) and "Выполненные работы по заказ-наряду" in c for c in row) for row in rows)
    has_materials_section = any(any(isinstance(c, str) and "Расходная накладная к заказ-наряду" in c for c in row) for row in rows)
    if not (has_labor_section or has_materials_section):
        return None

    meta = {
        "order_number": None,
        "order_date": None,
        "vehicle_make": None,
        "vehicle_model": None,
        "vehicle_vin": None,
        "vehicle_year": None,
    }
    labor_lines: list[dict] = []
    part_lines: list[dict] = []
    mode = None
    labor_cols: dict[str, int | None] = {}
    material_cols: dict[str, int | None] = {}
    labor_num_col: int | None = None
    material_num_col: int | None = None

    for row in rows:
        joined = " ".join(c for c in row if isinstance(c, str))

        if meta["order_number"] is None:
            m = _ORDER_RE.search(joined)
            if m:
                meta["order_number"] = m.group(1)
                meta["order_date"] = m.group(2)

        if meta["vehicle_make"] is None:
            m = _VEHICLE_RE.search(joined)
            if m:
                make_model = m.group(1).split(None, 1)
                meta["vehicle_make"] = make_model[0] if make_model else None
                meta["vehicle_model"] = make_model[1] if len(make_model) > 1 else None
                meta["vehicle_vin"] = m.group(3)
                year_m = _YEAR_RE.search(joined)
                if year_m:
                    meta["vehicle_year"] = int(year_m.group(1) or year_m.group(2))

        if "Выполненные работы по заказ-наряду" in joined:
            mode = "await_labor_header"
            continue
        if mode == "await_labor_header":
            num_col = _find_number_column(row)
            if num_col is not None:
                labor_num_col = num_col
                labor_cols = {
                    "catalog_code": _find_export_column(row, LABOR_CATALOG_CODE_ALIASES),
                    "description": _find_export_column(row, LABOR_DESCRIPTION_ALIASES),
                    "hourly_rate": _find_export_column(row, LABOR_HOURLY_RATE_ALIASES),
                    "norm_hours": _find_export_column(row, LABOR_NORM_HOURS_ALIASES),
                    "total": _find_export_column(row, LABOR_TOTAL_ALIASES),
                }
                if labor_cols["description"] is None:
                    logger.warning(
                        "parse_repair_order_export: не нашёл колонку с наименованием работы в заголовке %r — "
                        "раздел 'Выполненные работы' пропущен",
                        row,
                    )
                mode = "await_labor_index"
            continue
        if mode == "await_labor_index":
            # Строка-легенда номеров колонок под заголовком есть не во всех
            # выгрузках (см. _looks_like_data_row) — если её нет, эта же
            # строка уже данные, и её нельзя пропускать.
            mode = "labor"
            if _looks_like_data_row(row, labor_cols.get("description")):
                pass  # не continue — обработать эту же строку как данные ниже
            else:
                continue
        if mode == "labor":
            if "Итого работ" in joined:
                mode = None
                continue
            num_val = _clean(_export_cell(row, labor_num_col))
            if num_val and num_val.isdigit() and labor_cols.get("description") is not None:
                labor_lines.append(
                    {
                        "description": _clean(_export_cell(row, labor_cols["description"])),
                        "catalog_code": _clean(_export_cell(row, labor_cols["catalog_code"])),
                        "hourly_rate": _to_float(_export_cell(row, labor_cols["hourly_rate"])),
                        "norm_hours": _to_float(_export_cell(row, labor_cols["norm_hours"])),
                        "total": _to_float(_export_cell(row, labor_cols["total"])),
                    }
                )
            continue

        if "Расходная накладная к заказ-наряду" in joined:
            mode = "await_materials_header"
            continue
        if mode == "await_materials_header":
            num_col = _find_number_column(row)
            if num_col is not None:
                material_num_col = num_col
                material_cols = {
                    "article": _find_export_column(row, MATERIAL_ARTICLE_ALIASES),
                    "name": _find_export_column(row, MATERIAL_NAME_ALIASES),
                    "qty": _find_export_column(row, MATERIAL_QTY_ALIASES),
                    "price": _find_export_column(row, MATERIAL_PRICE_ALIASES),
                }
                if material_cols["name"] is None:
                    logger.warning(
                        "parse_repair_order_export: не нашёл колонку с наименованием запчасти в заголовке %r — "
                        "раздел 'Расходная накладная' пропущен",
                        row,
                    )
                mode = "await_materials_index"
            continue
        if mode == "await_materials_index":
            mode = "materials"
            if _looks_like_data_row(row, material_cols.get("name")):
                pass  # не continue — обработать эту же строку как данные ниже
            else:
                continue
        if mode == "materials":
            if any(
                marker in joined
                for marker in ("Итого по странице материалов", "Итого материалов", "Итого запчасти")
            ):
                mode = None
                continue
            num_val = _clean(_export_cell(row, material_num_col))
            if num_val and num_val.isdigit() and material_cols.get("name") is not None:
                part_lines.append(
                    {
                        "article": _clean(_export_cell(row, material_cols["article"])),
                        "name": _clean(_export_cell(row, material_cols["name"])),
                        "qty": _to_float(_export_cell(row, material_cols["qty"])),
                        "price": _to_float(_export_cell(row, material_cols["price"])),
                    }
                )
            continue

    if not labor_lines and not part_lines:
        return None

    return {"meta": meta, "labor_lines": labor_lines, "part_lines": part_lines}


def _normalize_brand_label(label: str) -> str:
    """label из ярлыка раздела/листа/названия марки в каталоге -> тот же
    вид, в котором марка обычно приходит из заказ-наряда (латиница) — иначе
    сравнение в matcher._contract_candidate_pool никогда не совпадёт
    (кириллица vs латиница — разные строки побайтово, даже после upper()).
    Составные ярлыки вида "GM (Шевроле, Опель)" или "ЛАДА Гранта" — берём
    первое узнанное название (из скобок или первое слово вне скобок);
    секция физически одна, а ContractPart.vehicle_make — одно значение,
    развести несколько марок из одной секции не на что.

    Справочник соответствий — таблица BrandAlias в БД (не константа в
    коде): заказчик работает не только с уже присланными файлами, новую
    марку/написание можно добавить через админку или файлом, не дожидаясь
    правки кода и пересборки (см. app/api/brand_aliases.py). Требует
    app_context — тот же приём, что и nomenclature_import.py, который так
    же обращается к моделям/db прямо из "чистого" парсера."""
    from app.extensions import db
    from app.models import BrandAlias

    label = label.strip()
    if not label:
        return label
    paren = re.search(r"\(([^)]+)\)", label)
    candidates = []
    if paren:
        candidates.extend(c.strip() for c in paren.group(1).split(","))
        candidates.append(label[: paren.start()].strip())
    else:
        candidates.append(label)

    for candidate in candidates:
        if not candidate:
            continue
        key = candidate.upper()
        match = BrandAlias.query.filter(
            db.func.upper(BrandAlias.alias) == key, BrandAlias.canonical_make.isnot(None)
        ).first()
        if match:
            return match.canonical_make
        first_word = key.split()[0] if key.split() else key
        if first_word != key:
            match = BrandAlias.query.filter(
                db.func.upper(BrandAlias.alias) == first_word, BrandAlias.canonical_make.isnot(None)
            ).first()
            if match:
                return match.canonical_make

    return (candidates[0] if candidates and candidates[0] else label).upper()


# Два реальных варианта заголовка листа-каталога у заказчика:
#   "Марка (модель) технического средства X[, Y и Z]" — список марок сразу
#   "Запчасти на автомобиль X Model"                   — одна марка+модель
# Дальше наверняка встретятся и другие формулировки того же смысла у других
# поставщиков — заказчик не ограничится присланными файлами.
_CATALOG_TITLE_MARKER = "Марка (модель) технического средства"
_CATALOG_TITLE_MARKER_SINGLE = "Запчасти на автомобиль"


def _sheet_title_row(ws) -> str:
    """Не только первая ячейка — у части файлов заказчика заголовок листа
    стоит не в колонке A (см. testdata/Приложение ГП10 №3 с падением
    75.xlsx, листы "Dewoo"/"газ": ('Запчасти на автомобиль Daewoo Nexia'
    лежит во ВТОРОЙ ячейке, первая пустая) — раньше это читало только
    first_row[0] и такие листы вообще не находились как каталог по маркам."""
    first_row = next(ws.iter_rows(min_row=1, max_row=1, values_only=True), None)
    if not first_row:
        return ""
    return " ".join(str(c) for c in first_row if c is not None and str(c).strip())


def parse_price_catalog_by_brand(file_path: str, vehicle_make: str | None) -> list[dict] | None:
    ext = os.path.splitext(file_path)[1].lower()
    if ext not in (".xlsx", ".xlsm"):
        return None

    import openpyxl

    wb = openpyxl.load_workbook(file_path, data_only=True, read_only=True)

    brand_sheets: dict[str, str] = {}
    for name in wb.sheetnames:
        ws = wb[name]
        title = _sheet_title_row(ws)

        if _CATALOG_TITLE_MARKER in title:
            remainder = title.replace(_CATALOG_TITLE_MARKER, "")
            brands = remainder.replace(",", " и ").split(" и ")
        elif _CATALOG_TITLE_MARKER_SINGLE in title:
            # "Запчасти на автомобиль 22172 Соболь" (реальный файл, лист
            # "газ") — марки в тексте нет вовсе, только модель ГАЗ. Модель
            # не по алиасу не распознать — в таком случае марку не тегируем
            # (вернётся None, лист всё равно попадёт в общий пул при
            # vehicle_make=None — лучше, чем совсем потерять данные листа).
            remainder = title.replace(_CATALOG_TITLE_MARKER_SINGLE, "")
            first_word = remainder.strip().split()[0] if remainder.strip() else ""
            if first_word and not first_word.isdigit():
                brands = [first_word]
            elif 0 < len(name) <= 20:
                # В тексте марки нет вовсе — только модель (реальный файл,
                # лист "газ": "Запчасти на автомобиль 22172 Соболь"). Имя
                # листа тут само по себе марка ("газ") — пробуем его как
                # запасной вариант, а не теряем весь лист.
                brands = [name]
            else:
                brands = []
        else:
            continue

        for brand in brands:
            # Реальный файл заказчика: "Марка (модель) технического средства
            # - ChevrolNiva ,ЛАДА Гранта, ..." — после маркера идёт тире, и
            # .strip() его не убирает (это не пробельный символ), так что
            # без .lstrip("-–— ") бренд сохранялся как "- CHEVROLNIVA" и
            # НИКОГДА не совпадал бы с реальной маркой из заказ-наряда
            # ("CHEVROLET") ни в одном сравнении — именно то самое "по Ниве
            # ничего не находит", о котором сообщал заказчик.
            brand = brand.strip().lstrip("-–—").strip()
            if brand:
                brand_sheets[_normalize_brand_label(brand)] = name

    if not brand_sheets:
        return None

    target = (vehicle_make or "").strip().upper()
    if target:
        matched_sheets = [brand_sheets[target]] if target in brand_sheets else []
        if not matched_sheets:
            logger.warning(
                "Марка %r не найдена в каталоге %s (доступны: %s) — сопоставление пойдёт только по расходным материалам",
                vehicle_make,
                file_path,
                sorted(brand_sheets),
            )
    else:
        # Марка не задана — файл разом по нескольким маркам (см. sheetnames
        # выше), и ни одна из них не "более правильная" по умолчанию.
        # Раньше в этом случае функция вообще не вызывалась (см. вызывающий
        # код) и разбор уходил в общий парсер одной таблицы, который эту
        # структуру (строка-заголовок с маркой + отдельная строка с
        # колонками) не понимает и падает с "не удалось найти колонку с
        # наименованием" — реальный собранный руками файл заказчика именно
        # так и падал. Раз колонка "Марка" у ContractPart всё равно не
        # хранится (запчасти различаются по артикулу, а не по марке), для
        # запчастей безопасно и правильно просто взять ВСЕ найденные листы.
        seen_sheets: set[str] = set()
        matched_sheets = []
        for sheet_name in brand_sheets.values():
            if sheet_name not in seen_sheets:
                seen_sheets.add(sheet_name)
                matched_sheets.append(sheet_name)

    # Лист -> марка(и), которые на нём объявлены (обратный словарь к
    # brand_sheets) — один лист может быть на несколько марок сразу
    # ("Hyundai/Kia"), тогда строки помечаем первой из них: разделять один
    # физический прайс-лист на несколько идентичных копий ради тега не имеет
    # смысла, а для фильтрации в matcher._contract_candidate_pool важно лишь
    # не перепутать лист одной марки с листом другой.
    sheet_brand: dict[str, str] = {}
    for brand, sheet_name in brand_sheets.items():
        sheet_brand.setdefault(sheet_name, brand)

    results: list[dict] = []

    for matched_sheet in matched_sheets:
        ws = wb[matched_sheet]
        row_brand = sheet_brand.get(matched_sheet)

        # И порядок колонок, и то, на какой именно строке шапка, отличаются
        # у разных поставщиков даже внутри одного файла (testdata/Приложение
        # ГП10 №3 с падением 75.xlsx): у Nissan/Toyota шапка колонок —
        # отдельная строка 2 (№|Наименование|Цена|Характеристики). А у
        # Dewoo/газ заголовок листа и подписи колонок вообще слиты в одну
        # строку 1 ('Запчасти на автомобиль 22172 Соболь', 'ед.изм', 'цена')
        # — при этом "Наименование" там нигде НЕ подписано текстом вообще
        # (колонка 1 подразумевается по умолчанию), а "цена" — подписана.
        # Поэтому наименование/цену/артикул ищем НЕЗАВИСИМО друг от друга по
        # первым трём строкам (не требуя, чтобы оба нашлись в ОДНОЙ и той же
        # строке) — то, что не удалось найти по алиасу, остаётся на старой
        # позиции (1=наименование, 2=цена, 3=артикул), под которую формат
        # изначально и был написан. Данные начинаются сразу после самой
        # последней из строк, где хоть что-то нашлось (не путать заголовок
        # с первой строкой данных).
        name_col, price_col, article_col = 1, 2, 3
        detected_name_col = detected_price_col = detected_article_col = None
        last_detected_row: int | None = None
        for header_idx, row in enumerate(ws.iter_rows(min_row=1, max_row=3, values_only=True), start=1):
            row = list(row)
            found_name = _find_export_column(row, NAME_COLUMN_ALIASES)
            found_price = _find_export_column(row, PRICE_COLUMN_ALIASES)
            found_article = _find_export_column(row, ARTICLE_COLUMN_ALIASES)
            if found_name is not None:
                detected_name_col = found_name
                last_detected_row = header_idx
            if found_price is not None:
                detected_price_col = found_price
                last_detected_row = header_idx
            if found_article is not None:
                detected_article_col = found_article
                last_detected_row = header_idx
        if detected_name_col is not None:
            name_col = detected_name_col
        if detected_price_col is not None:
            price_col = detected_price_col
        if detected_article_col is not None:
            article_col = detected_article_col
        elif article_col in (name_col, price_col):
            # Позиция артикула по умолчанию (3) кем-то уже занята из-за
            # обнаруженных наименования/цены на других позициях (реальный
            # случай — Dewoo/газ: цена нашлась на позиции 3, своей отдельной
            # колонки под артикул в этих листах вообще нет) — читать оттуда
            # же под видом артикула значило бы задвоить цену в оба поля.
            article_col = None
        # Ничего не нашли по алиасам вовсе — старое допущение (шапка на
        # строке 2, данные с 3-й), под которое формат изначально написан;
        # нашли хоть что-то — данные сразу после самой поздней из строк,
        # где что-то обнаружилось (не перепутать шапку со строкой данных).
        data_start_row = (last_detected_row + 1) if last_detected_row is not None else 3

        for row in ws.iter_rows(min_row=data_start_row, values_only=True):
            name = _export_cell(list(row), name_col)
            name = _clean(name)
            if not name:
                continue
            results.append(
                {
                    "article": _clean(_export_cell(list(row), article_col)),
                    "name": name,
                    "qty": None,
                    "price": _to_float(_export_cell(list(row), price_col)),
                    "vehicle_make": row_brand,
                }
            )

    if "Расходные материалы" in wb.sheetnames:
        ws = wb["Расходные материалы"]
        for row in ws.iter_rows(min_row=3, values_only=True):
            name = _clean(row[1]) if len(row) > 1 else None
            if not name:
                continue
            results.append(
                {
                    "article": None,
                    "name": name,
                    "qty": None,
                    "price": _to_float(row[3]) if len(row) > 3 else None,
                    # Расходники общие для всех марок — не привязываем.
                    "vehicle_make": None,
                }
            )

    return results


# Раздел в single-sheet каталоге (см. parse_price_catalog_single_sheet_sections)
# может быть не маркой, а общей категорией "для всех марок" — тогда её не
# нужно превращать в фиктивную "марку", по которой заказ-наряд конкретного
# автомобиля никогда не совпадёт (тот же принцип, что "Расходные материалы"
# у parse_price_catalog_by_brand выше).
_SECTION_LABEL_IS_UNIVERSAL_RE = re.compile(
    r"расходн|материал|неоригинальн|масл|жидкост", re.IGNORECASE
)


def parse_price_catalog_single_sheet_sections(file_path: str) -> list[dict] | None:
    """Каталог ОДНИМ листом на несколько марок сразу, где раздел на марку —
    не отдельный лист (см. parse_price_catalog_by_brand выше), а
    строка-маркер внутри самого листа: единственная непустая ячейка
    (название марки), дальше обычные строки данных до следующего маркера.

    Реальный файл заказчика (testdata/Приложение №1 ИП Даянова З.Р..xlsx) —
    один лист, 25000+ строк, разделы LADA (ВАЗ)/УАЗ/ГАЗ/ПАЗ/TOYOTA/
    GM (Шевроле, Опель)/Неоригинальные запчасти/Масла... — шапка колонок
    ("№ п/п | Артикул производителя | Наименование ... | Ед. изм. | Цена...")
    вдобавок не в первой строке листа, а в третьей (после названия
    приложения и текстового заголовка раздела), поэтому обычный
    _dataframe_to_lines (шапка = первая строка pandas) тут в принципе не
    мог сработать.

    Возвращает None, если ни на одном листе не нашлась строка-шапка с
    узнаваемой колонкой "наименование" — тогда вызывающий код пробует
    другие парсеры (см. contract_catalog_import.py)."""
    ext = os.path.splitext(file_path)[1].lower()
    if ext not in (".xlsx", ".xlsm"):
        return None

    import openpyxl

    wb = openpyxl.load_workbook(file_path, data_only=True, read_only=True)

    results: list[dict] = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        header_row_idx = None
        cols = None
        for i, row in enumerate(ws.iter_rows(min_row=1, max_row=15, values_only=True), start=1):
            row = list(row)
            name_col = _find_export_column(row, NAME_COLUMN_ALIASES)
            price_col = _find_export_column(row, PRICE_COLUMN_ALIASES)
            if name_col is not None and price_col is not None:
                header_row_idx = i
                cols = {
                    "article": _find_export_column(row, ARTICLE_COLUMN_ALIASES),
                    "name": name_col,
                    "price": price_col,
                }
                break
        if header_row_idx is None:
            continue

        current_brand: str | None = None
        for row in ws.iter_rows(min_row=header_row_idx + 1, values_only=True):
            row = list(row)
            name = _export_cell(row, cols["name"])
            price = _export_cell(row, cols["price"])
            has_name = name is not None and str(name).strip()
            has_price = price is not None and str(price).strip() != ""
            if not (has_name and has_price):
                # Строка-маркер раздела (или пустая строка-разделитель) —
                # берём первую непустую ячейку как название марки.
                label = next((c for c in row if c is not None and str(c).strip()), None)
                if label is not None:
                    label = str(label).strip()
                    current_brand = (
                        None if _SECTION_LABEL_IS_UNIVERSAL_RE.search(label) else _normalize_brand_label(label)
                    )
                continue
            results.append(
                {
                    "article": _clean(_export_cell(row, cols["article"])),
                    "name": _clean(name),
                    "qty": None,
                    "price": _to_float(price),
                    "vehicle_make": current_brand,
                }
            )

    return results or None
