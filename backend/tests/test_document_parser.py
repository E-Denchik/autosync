import os

import openpyxl
import pandas as pd
import pytest

from app.services.document_parser import (
    DocumentParseError,
    _normalize_brand_label,
    parse_document,
    parse_document_with_ocr_fallback,
    parse_hourly_rate_table,
    parse_price_catalog_by_brand,
    parse_price_catalog_single_sheet_sections,
    parse_repair_order_export,
)

TESTDATA_DIR = os.path.join(os.path.dirname(__file__), "..", "..", "testdata")

ROWS = [
    {"Артикул": "ABC-1", "Наименование": "Тормозной диск", "Кол-во": 2, "Цена": "1 500,50"},
    {"Артикул": "", "Наименование": "Фильтр масляный", "Кол-во": 1, "Цена": 350},
]


def _expected():
    return [
        {"article": "ABC-1", "name": "Тормозной диск", "qty": 2.0, "price": 1500.50},
        {"article": None, "name": "Фильтр масляный", "qty": 1.0, "price": 350.0},
    ]


def test_parse_xlsx(tmp_path):
    path = tmp_path / "contract.xlsx"
    pd.DataFrame(ROWS).to_excel(path, index=False, engine="openpyxl")
    assert parse_document(str(path)) == _expected()


def test_parse_csv_comma_utf8(tmp_path):
    path = tmp_path / "contract.csv"
    pd.DataFrame(ROWS).to_csv(path, index=False, encoding="utf-8-sig")
    assert parse_document(str(path)) == _expected()


def test_parse_csv_semicolon_cp1251(tmp_path):
    path = tmp_path / "contract.csv"
    pd.DataFrame(ROWS).to_csv(path, index=False, sep=";", encoding="cp1251")
    assert parse_document(str(path)) == _expected()


def test_parse_csv_semicolon_with_comma_inside_header_and_cells(tmp_path):
    """Регрессия: pd.read_csv(sep=None) автоопределяет разделитель по
    небольшой выборке строк и путается, если запятая встречается прямо в
    данных — например, в заголовке "Цена, руб." или в ячейке со списком
    марок через запятую ("Renault Sandero, Nissan Almera, ...", реальный
    случай из testdata/Нормочасы.csv — 8 строк, часть с несколькими марками
    через запятую в одной ячейке). Раньше это падало с "Expected N fields...
    saw M", потому что sep=None на части файла решал, будто разделитель —
    запятая, хотя весь файл на самом деле разделён ';'."""
    path = os.path.join(TESTDATA_DIR, "Нормочасы.csv")
    lines = parse_hourly_rate_table(path)

    assert len(lines) == 15  # одна ячейка "Renault Sandero, Nissan Almera, ..." разворачивается в 5 строк
    rates = {(line["vehicle_make"], line["vehicle_model"]): line["hourly_rate"] for line in lines}
    assert rates[("Chevrolet", "Niva")] == 540.0
    assert rates[("Hyundai", "Accent")] == 720.0


def test_parse_ods(tmp_path):
    path = tmp_path / "contract.ods"
    pd.DataFrame(ROWS).to_excel(path, index=False, engine="odf")
    assert parse_document(str(path)) == _expected()


def test_parse_docx(tmp_path):
    from docx import Document

    document = Document()
    table = document.add_table(rows=1, cols=4)
    for cell, header in zip(table.rows[0].cells, ["Артикул", "Наименование", "Кол-во", "Цена"]):
        cell.text = header
    for row_data in ROWS:
        cells = table.add_row().cells
        for cell, key in zip(cells, ["Артикул", "Наименование", "Кол-во", "Цена"]):
            cell.text = str(row_data[key])
    path = tmp_path / "order.docx"
    document.save(path)

    assert parse_document(str(path)) == _expected()


def test_ocr_fallback_not_needed_llm_client_untouched_for_normal_xlsx(tmp_path):
    """Regression guard: обычный, распознаваемый жёсткими алиасами файл не
    должен вообще трогать LLM — фоллбэк только для того, что жёсткий разбор
    не осилил."""
    from unittest.mock import MagicMock

    path = tmp_path / "contract.xlsx"
    pd.DataFrame(ROWS).to_excel(path, index=False, engine="openpyxl")

    llm_client = MagicMock()
    result = parse_document_with_ocr_fallback(str(path), llm_client, ["article", "name", "qty", "price"])

    assert result == _expected()
    llm_client.extract_table_from_text.assert_not_called()


def test_ocr_fallback_kicks_in_when_xlsx_headers_are_not_recognized(tmp_path):
    """Заказчик: "чтобы файлы с любым содержимым считывались" — раньше
    нестандартная шапка (не совпавшая ни с одним алиасом в
    NAME_COLUMN_ALIASES) сразу роняла DocumentParseError, даже если файл
    вполне читаемый xlsx с данными внутри. Теперь вместо немедленного отказа
    сырой текст файла (df.to_csv) уходит на LLM-извлечение — тем же путём,
    что уже был для сканов/фото."""
    from unittest.mock import MagicMock

    path = tmp_path / "weird.xlsx"
    pd.DataFrame(
        [{"Поз.": "ABC-1", "Что за деталь": "Тормозной диск", "N": 2, "Руб.": 1500.5}]
    ).to_excel(path, index=False, engine="openpyxl")

    with pytest.raises(DocumentParseError):
        parse_document(str(path))  # жёсткий разбор действительно не справляется — база теста верна

    llm_client = MagicMock()
    llm_client.extract_table_from_text.return_value = [
        {"article": "ABC-1", "name": "Тормозной диск", "qty": 2, "price": 1500.5}
    ]

    result = parse_document_with_ocr_fallback(str(path), llm_client, ["article", "name", "qty", "price"])

    assert result == [{"article": "ABC-1", "name": "Тормозной диск", "qty": 2.0, "price": 1500.5}]
    passed_text = llm_client.extract_table_from_text.call_args[0][0]
    assert "ABC-1" in passed_text and "Тормозной диск" in passed_text


def test_ocr_fallback_raises_original_error_when_raw_text_extraction_also_empty(tmp_path):
    """Если и сырой текст достать не из чего (документ и правда пустой —
    ни таблиц, ни абзацев) — показываем исходную, более конкретную ошибку
    жёсткого парсера, а не глотаем её ради LLM-фоллбэка, которому нечего
    было бы передать."""
    from docx import Document
    from unittest.mock import MagicMock

    path = tmp_path / "empty.docx"
    Document().save(path)  # ни абзацев с текстом, ни таблиц

    llm_client = MagicMock()
    with pytest.raises(DocumentParseError, match="таблиц"):
        parse_document_with_ocr_fallback(str(path), llm_client, ["article", "name", "qty", "price"])
    llm_client.extract_table_from_text.assert_not_called()


def test_docx_with_one_bad_table_and_one_good_table_keeps_the_good_one(tmp_path):
    """Regression: docx с несколькими таблицами (нужная ведомость +,
    например, служебная таблица реквизитов) раньше падал целиком, если
    ХОТЯ БЫ одна таблица не имела узнаваемой колонки "наименование" — даже
    если нужная таблица была прочитана бы нормально."""
    from docx import Document

    document = Document()

    # "Служебная" таблица без узнаваемых колонок.
    junk = document.add_table(rows=1, cols=2)
    junk.rows[0].cells[0].text = "ИНН"
    junk.rows[0].cells[1].text = "770000000"

    good = document.add_table(rows=1, cols=4)
    for cell, header in zip(good.rows[0].cells, ["Артикул", "Наименование", "Кол-во", "Цена"]):
        cell.text = header
    for row_data in ROWS:
        cells = good.add_row().cells
        for cell, key in zip(cells, ["Артикул", "Наименование", "Кол-во", "Цена"]):
            cell.text = str(row_data[key])

    path = tmp_path / "mixed.docx"
    document.save(path)

    assert parse_document(str(path)) == _expected()


def test_unsupported_extension_raises(tmp_path):
    path = tmp_path / "contract.txt"
    path.write_text("что угодно")
    with pytest.raises(DocumentParseError):
        parse_document(str(path))


def _write_repair_order_export(path, labor_header: dict, labor_rows: list[dict]):
    """Строит минимальную печатную форму заказ-наряда 1С (только раздел
    "Выполненные работы") с ПРОИЗВОЛЬНЫМ порядком колонок — labor_header и
    каждая строка labor_rows задаются словарём {индекс_колонки: значение},
    индексы 0-based, как их читает pandas (см. parse_repair_order_export)."""
    wb = openpyxl.Workbook()
    ws = wb.active

    def put(row_idx: int, values: dict):
        for col0, value in values.items():
            ws.cell(row=row_idx, column=col0 + 1, value=value)

    put(1, {2: "Заказ-наряд №  0001 от 01.01.2026"})
    put(3, {2: "Выполненные работы по заказ-наряду"})
    put(5, labor_header)
    put(6, {i: str(i) for i in range(1, 9)})  # строка-легенда "1 2 3 ... 8", как в реальном файле
    row_idx = 7
    for i, data in enumerate(labor_rows, start=1):
        put(row_idx, {1: str(i), **data})
        row_idx += 1
    put(row_idx, {1: "Итого работ:"})

    wb.save(path)


STANDARD_LABOR_HEADER = {
    1: "№",
    2: "№ кат.",
    3: "Наименование",
    8: "Кол. оп.",
    9: "Цена н/ч",
    10: "Норма",
    11: "н/ч",
    12: "Всего",
}


def test_parse_repair_order_export_reads_labor_columns_by_header_not_fixed_position(tmp_path):
    path = tmp_path / "order.xlsx"
    _write_repair_order_export(
        path,
        STANDARD_LABOR_HEADER,
        [{3: "ДВС снятие", 9: 1000, 10: 28, 12: 28000}],
    )
    result = parse_repair_order_export(str(path))
    assert result["labor_lines"] == [
        {"description": "ДВС снятие", "catalog_code": None, "hourly_rate": 1000.0, "norm_hours": 28.0, "total": 28000.0}
    ]


def test_parse_repair_order_export_adapts_to_reordered_columns(tmp_path):
    """Регрессия: раньше колонки читались строго по фиксированной позиции
    (row[9]/row[10]/row[12]) — совпадало с ОДНИМ конкретным шаблоном отчёта
    1С, но другая конфигурация вполне может расставить колонки иначе.
    Здесь порядок совсем другой (наименование раньше, норма и цена
    поменяны местами, сумма в дальней колонке) — результат должен быть
    таким же, потому что колонки ищутся по заголовку, а не по индексу."""
    path = tmp_path / "order.xlsx"
    reordered_header = {
        1: "№",
        2: "Наименование работы",
        5: "№ кат.",
        9: "Норма часов",
        10: "Цена нормо-часа",
        13: "Сумма",
    }
    _write_repair_order_export(
        path,
        reordered_header,
        [{2: "ДВС снятие", 5: "K-1", 9: 28, 10: 1000, 13: 28000}],
    )
    result = parse_repair_order_export(str(path))
    assert result["labor_lines"] == [
        {"description": "ДВС снятие", "catalog_code": "K-1", "hourly_rate": 1000.0, "norm_hours": 28.0, "total": 28000.0}
    ]


def test_parse_repair_order_export_skips_labor_section_when_description_column_not_found(tmp_path):
    """Если заголовок совсем не похож на ожидаемый (колонки не нашлись) —
    раздел просто пропускается, а не заполняется мусором из случайных
    колонок (тихая порча данных хуже отсутствия данных)."""
    path = tmp_path / "order.xlsx"
    unrecognizable_header = {1: "№", 2: "Column A", 3: "Column B"}
    _write_repair_order_export(path, unrecognizable_header, [{2: "что-то", 3: "что-то ещё"}])
    result = parse_repair_order_export(str(path))
    assert result is None


def _write_repair_order_materials_export(path, materials_header: dict, materials_rows: list[dict]):
    wb = openpyxl.Workbook()
    ws = wb.active

    def put(row_idx: int, values: dict):
        for col0, value in values.items():
            ws.cell(row=row_idx, column=col0 + 1, value=value)

    put(1, {2: "Заказ-наряд №  0001 от 01.01.2026"})
    put(3, {2: "Расходная накладная к заказ-наряду"})
    put(5, materials_header)
    put(6, {i: str(i) for i in range(1, 8)})
    row_idx = 7
    for i, data in enumerate(materials_rows, start=1):
        put(row_idx, {1: str(i), **data})
        row_idx += 1
    put(row_idx, {1: "Итого материалов:"})

    wb.save(path)


def test_parse_repair_order_export_adapts_to_reordered_material_columns(tmp_path):
    path = tmp_path / "order.xlsx"
    reordered_header = {1: "№", 2: "Название", 4: "Артикул", 9: "Количество", 11: "Цена"}
    _write_repair_order_materials_export(
        path,
        reordered_header,
        [{2: "Прокладка головки блока цилиндров", 4: "2231125013", 9: 1, 11: 3100}],
    )
    result = parse_repair_order_export(str(path))
    assert result["part_lines"] == [
        {"article": "2231125013", "name": "Прокладка головки блока цилиндров", "qty": 1.0, "price": 3100.0}
    ]


def test_parse_repair_order_export_handles_layout_without_leading_blank_column(tmp_path):
    """Реальный файл заказчика (testdata/repair_order_1_final.xlsx) устроен
    иначе, чем печатная форма 1С, под которую был написан парсер изначально:
    - "№" сразу в колонке A (не в B — раньше жёстко бралась row[1]);
    - данные идут сразу после заголовка, без строки-легенды "1 2 3 ... 9";
    - "Автомобиль: MAKE MODEL VIN: ... YYYY г." без "гос. номер:" и без
      "год вып.";
    - раздел запчастей завершается "Итого запчасти:", а не "Итого
      материалов:".
    Раньше это приводило к тому, что parse_repair_order_export не находил
    вообще ничего и возвращал None, а загрузка падала с "не удалось найти
    колонку с наименованием" в общем однотабличном парсере."""
    path = os.path.join(TESTDATA_DIR, "repair_order_1_final.xlsx")
    result = parse_repair_order_export(path)

    assert result is not None
    assert result["meta"]["order_number"] == "1"
    assert result["meta"]["order_date"] == "17.08.2026"
    assert result["meta"]["vehicle_make"] == "HYUNDAI"
    assert result["meta"]["vehicle_model"] == "IX35"
    assert result["meta"]["vehicle_vin"] == "TMAJU81BCCJ238125"
    assert result["meta"]["vehicle_year"] == 2011

    assert len(result["labor_lines"]) == 3
    assert result["labor_lines"][0] == {
        "description": "ДВС снятие",
        "catalog_code": None,
        "hourly_rate": 1170.0,
        "norm_hours": 28.0,
        "total": 32760.0,
    }

    assert len(result["part_lines"]) == 3
    names = [p["name"] for p in result["part_lines"]]
    assert "Расходные материалы" in names
    assert any("поршень с кольцами" in n.lower() for n in names)


def test_parse_hourly_rate_table_xlsx(tmp_path):
    path = tmp_path / "rates.xlsx"
    pd.DataFrame(
        [{"Марка": "Hyundai", "Ставка, руб/ч": 800}, {"Марка": "Toyota", "Ставка, руб/ч": 900}]
    ).to_excel(path, index=False, engine="openpyxl")
    assert parse_hourly_rate_table(str(path)) == [
        {"vehicle_make": "Hyundai", "vehicle_model": None, "hourly_rate": 800.0},
        {"vehicle_make": "Toyota", "vehicle_model": None, "hourly_rate": 900.0},
    ]


def test_parse_hourly_rate_table_adapts_to_different_column_names(tmp_path):
    """Образец файла заказчика заранее неизвестен — колонки называются как
    угодно ("Brand"/"Цена нормо-часа"), ищутся по синонимам, а не по одному
    жёстко заданному имени."""
    path = tmp_path / "rates.xlsx"
    pd.DataFrame(
        [{"Brand": "KIA", "Цена нормо-часа": "1 200,50"}]
    ).to_excel(path, index=False, engine="openpyxl")
    assert parse_hourly_rate_table(str(path)) == [{"vehicle_make": "KIA", "vehicle_model": None, "hourly_rate": 1200.50}]


def test_parse_hourly_rate_table_csv(tmp_path):
    path = tmp_path / "rates.csv"
    pd.DataFrame([{"Марка ТС": "Hyundai", "Стоимость": 800}]).to_csv(path, index=False, encoding="utf-8-sig")
    assert parse_hourly_rate_table(str(path)) == [{"vehicle_make": "Hyundai", "vehicle_model": None, "hourly_rate": 800.0}]


def test_parse_hourly_rate_table_skips_blank_make_and_non_positive_rate(tmp_path):
    path = tmp_path / "rates.xlsx"
    pd.DataFrame(
        [
            {"Марка": "Hyundai", "Ставка": 800},
            {"Марка": "", "Ставка": 900},
            {"Марка": "Toyota", "Ставка": 0},
            {"Марка": "Kia", "Ставка": -100},
        ]
    ).to_excel(path, index=False, engine="openpyxl")
    assert parse_hourly_rate_table(str(path)) == [{"vehicle_make": "Hyundai", "vehicle_model": None, "hourly_rate": 800.0}]


def test_parse_hourly_rate_table_raises_when_make_column_not_found(tmp_path):
    path = tmp_path / "rates.xlsx"
    pd.DataFrame([{"Column A": "x", "Ставка": 800}]).to_excel(path, index=False, engine="openpyxl")
    with pytest.raises(DocumentParseError):
        parse_hourly_rate_table(str(path))


def test_parse_hourly_rate_table_raises_when_rate_column_not_found(tmp_path):
    path = tmp_path / "rates.xlsx"
    pd.DataFrame([{"Марка": "Hyundai", "Column B": 800}]).to_excel(path, index=False, engine="openpyxl")
    with pytest.raises(DocumentParseError):
        parse_hourly_rate_table(str(path))


def test_parse_hourly_rate_table_docx_real_customer_file():
    """Реальный файл заказчика — приложение к тендерному контракту:
    таблица "Марка (модель) | Цена единицы услуги" внутри Word-документа,
    марка+модель одной ячейкой, несколько моделей через запятую на одну
    цену, плюс итоговая строка "ИТОГО с учетом аукционного снижения", её
    не должно быть в результате."""
    path = os.path.join(TESTDATA_DIR, "Нормочасы.docx")
    result = parse_hourly_rate_table(path)

    assert len(result) == 15
    assert not any("итого" in r["vehicle_make"].lower() for r in result)

    by_model = {(r["vehicle_make"], r["vehicle_model"]): r["hourly_rate"] for r in result}
    assert by_model[("Chevrolet", "Niva")] == 540.0
    assert by_model[("Hyundai", "Accent")] == 720.0
    assert by_model[("Hyundai", "Sonata")] == 720.0
    assert by_model[("Hyundai", "Tucson")] == 810.0
    assert by_model[("Hyundai", "IX35")] == 810.0
    assert by_model[("Hyundai", "Santa Fe")] == 810.0
    assert by_model[("Audi", "A8")] == 900.0


def test_parse_hourly_rate_table_splits_comma_separated_make_model_pairs(tmp_path):
    path = tmp_path / "rates.xlsx"
    pd.DataFrame([{"Марка (модель)": "Renault Sandero, Nissan Teana, Hyundai Accent", "Цена": 720}]).to_excel(
        path, index=False, engine="openpyxl"
    )
    result = parse_hourly_rate_table(str(path))
    assert result == [
        {"vehicle_make": "Renault", "vehicle_model": "Sandero", "hourly_rate": 720.0},
        {"vehicle_make": "Nissan", "vehicle_model": "Teana", "hourly_rate": 720.0},
        {"vehicle_make": "Hyundai", "vehicle_model": "Accent", "hourly_rate": 720.0},
    ]


def test_parse_hourly_rate_table_uses_separate_model_column_when_present(tmp_path):
    path = tmp_path / "rates.xlsx"
    pd.DataFrame([{"Марка": "Hyundai", "Модель": "IX35", "Ставка": 810}]).to_excel(
        path, index=False, engine="openpyxl"
    )
    assert parse_hourly_rate_table(str(path)) == [
        {"vehicle_make": "Hyundai", "vehicle_model": "IX35", "hourly_rate": 810.0}
    ]


def test_parse_hourly_rate_table_bare_make_has_no_model(tmp_path):
    path = tmp_path / "rates.xlsx"
    pd.DataFrame([{"Марка": "Hyundai", "Ставка": 800}]).to_excel(path, index=False, engine="openpyxl")
    assert parse_hourly_rate_table(str(path)) == [
        {"vehicle_make": "Hyundai", "vehicle_model": None, "hourly_rate": 800.0}
    ]


def test_parse_hourly_rate_table_excludes_total_row(tmp_path):
    path = tmp_path / "rates.xlsx"
    pd.DataFrame(
        [
            {"Марка": "Hyundai", "Ставка": 800},
            {"Марка": "ИТОГО с учетом аукционного снижения (55%):", "Ставка": 6209.99},
        ]
    ).to_excel(path, index=False, engine="openpyxl")
    result = parse_hourly_rate_table(str(path))
    assert result == [{"vehicle_make": "Hyundai", "vehicle_model": None, "hourly_rate": 800.0}]


def test_parse_hourly_rate_table_unsupported_extension_raises(tmp_path):
    path = tmp_path / "rates.zip"
    path.write_text("что угодно")
    with pytest.raises(DocumentParseError):
        parse_hourly_rate_table(str(path))


# Тендерное приложение — бумажный документ по своей природе (печать,
# подпись, печать организации), фото/скан вместо цифрового оригинала —
# не гипотетический случай, а то же самое, что уже поддержано для
# заказ-нарядов/договоров (см. services/ocr.py). Ниже — та же реальная
# картина, что и docx/pdf-тесты выше, но пришедшая как фото: реальный OCR-
# текст с этого же файла (снят один раз через настоящий Tesseract на
# рендере реального Нормочасы.docx — см. историю разработки), а не выдуманный.
_REAL_OCR_TEXT_FROM_CUSTOMER_RATE_TABLE = """\
Идентификационный код закупки: 262631610547363160100100010014520244

к Контракту № 02-2026/ЭА от 02.03.2026

ПЕРЕЧЕНЬ ЦЕН ЕДИНИЦ УСЛУГ

№ Марка (модель)
п/п Цена единицы услуги, руб.
Chevrolet Niva 540
Renault Sandero, Nissan Almera Classik, Nissan Teana, 720

Hyundai Accent, Hyundai Sonata

KIA Spjrtage 810

Renault Duster 810

Volkswagen Multivan, Ford Tranzit 810
Hyundai Tucson, Hyundai 1X35, Hyundai Santa Fe 810
Toyota Land Cruiser 810

Audi A8 900

ИТОГО с учетом аукционного снижения (55%): 6 209.99999972623
"""


def test_parse_hourly_rate_table_image_uses_ocr_and_llm_extraction(tmp_path, monkeypatch):
    from unittest.mock import MagicMock

    from PIL import Image

    img_path = tmp_path / "rates.png"
    Image.new("RGB", (10, 10)).save(img_path)
    monkeypatch.setattr(
        "pytesseract.image_to_string", lambda image, lang=None: _REAL_OCR_TEXT_FROM_CUSTOMER_RATE_TABLE
    )

    llm_client = MagicMock()
    llm_client.extract_table_from_text.return_value = [
        {"vehicle_make": "Chevrolet Niva", "vehicle_model": None, "hourly_rate": 540},
        {
            "vehicle_make": "Renault Sandero, Nissan Almera Classik, Nissan Teana, Hyundai Accent, Hyundai Sonata",
            "vehicle_model": None,
            "hourly_rate": 720,
        },
        {"vehicle_make": "Hyundai Tucson, Hyundai IX35, Hyundai Santa Fe", "vehicle_model": None, "hourly_rate": 810},
        {"vehicle_make": "ИТОГО с учетом аукционного снижения (55%):", "vehicle_model": None, "hourly_rate": 6209.99},
    ]

    result = parse_hourly_rate_table(str(img_path), llm_client=llm_client)

    by_model = {(r["vehicle_make"], r["vehicle_model"]): r["hourly_rate"] for r in result}
    assert by_model[("Hyundai", "IX35")] == 810.0
    assert by_model[("Hyundai", "Accent")] == 720.0
    assert by_model[("Chevrolet", "Niva")] == 540.0
    assert not any("итого" in make.lower() for make, _ in by_model)

    # OCR-текст реально дошёл до LLM, а не какая-то заглушка.
    passed_text = llm_client.extract_table_from_text.call_args[0][0]
    assert "Chevrolet Niva" in passed_text


def test_parse_hourly_rate_table_image_without_llm_client_raises_clear_error(tmp_path, monkeypatch):
    from PIL import Image

    img_path = tmp_path / "rates.png"
    Image.new("RGB", (10, 10)).save(img_path)
    monkeypatch.setattr("pytesseract.image_to_string", lambda image, lang=None: "неважно")

    with pytest.raises(DocumentParseError):
        parse_hourly_rate_table(str(img_path))


def test_parse_hourly_rate_table_pdf_without_text_layer_falls_back_to_ocr(tmp_path, monkeypatch):
    """PDF без текстового слоя (тоже фото/скан, просто упакованное в PDF) —
    extract_pdf_tables() ничего не находит, должен сработать тот же OCR-путь,
    что и для картинки."""
    from unittest.mock import MagicMock

    path = tmp_path / "rates.pdf"
    path.write_bytes(b"%PDF-1.4\n%%EOF\n")  # заведомо без извлекаемых таблиц

    monkeypatch.setattr("app.services.document_parser.extract_pdf_tables", lambda p: [])

    from app.services import ocr as ocr_module

    monkeypatch.setattr(ocr_module, "extract_text", lambda p: "Hyundai IX35 810\n")

    llm_client = MagicMock()
    llm_client.extract_table_from_text.return_value = [
        {"vehicle_make": "Hyundai", "vehicle_model": "IX35", "hourly_rate": 810}
    ]

    result = parse_hourly_rate_table(str(path), llm_client=llm_client)
    assert result == [{"vehicle_make": "Hyundai", "vehicle_model": "IX35", "hourly_rate": 810.0}]


def test_normalize_brand_label_translits_cyrillic_to_latin(app):
    """Каталоги заказчика пишут марку то латиницей (как в заказ-наряде),
    то кириллицей — без транслитерации сравнение в
    matcher._contract_candidate_pool никогда бы не совпало. Справочник —
    таблица BrandAlias в БД (засеяна миграцией из старого хардкода), не
    константа в коде — поэтому вызов требует app_context."""
    with app.app_context():
        assert _normalize_brand_label("Шевроле") == "CHEVROLET"
        assert _normalize_brand_label("TOYOTA") == "TOYOTA"  # уже латиница — не трогаем


def test_normalize_brand_label_extracts_from_parentheses(app):
    """Реальный ярлык раздела каталога заказчика: "GM (Шевроле, Опель)" —
    берём первое узнанное название из скобок, а не бесполезную аббревиатуру
    концерна снаружи."""
    with app.app_context():
        assert _normalize_brand_label("GM (Шевроле, Опель)") == "CHEVROLET"


def test_normalize_brand_label_takes_first_word_of_brand_plus_model(app):
    """"ЛАДА Гранта"/"ВАЗ 2131" — марка+модель одной строкой (нет отдельной
    колонки модели в этом формате) — для ContractPart.vehicle_make нужна
    только марка, и разные модели одной марки должны сходиться к одному
    значению, а не плодить несовпадающие варианты."""
    with app.app_context():
        assert _normalize_brand_label("ЛАДА Гранта") == "LADA"
        assert _normalize_brand_label("ВАЗ 2131") == "LADA"
        assert _normalize_brand_label("Лада Самара") == "LADA"


def test_parse_price_catalog_by_brand_strips_leading_dash_from_brand_token(app):
    """Регрессия по реальному файлу заказчика (testdata/Приложение ГП10 №3
    с падением 75.xlsx, лист "Ваз,Нива"): заголовок листа — "Марка (модель)
    технического средства - ChevrolNiva ,ЛАДА Гранта, ...", и до фикса
    первый бренд-токен после маркера сохранялся как "- CHEVROLNIVA" (тире
    не пробельный символ, .strip() его не убирал) — такой ключ никогда не
    совпадал бы с маркой "CHEVROLET" из реального заказ-наряда. Это и есть
    техническая причина жалобы заказчика "по Ниве ничего не находит"."""
    path = os.path.join(TESTDATA_DIR, "Приложение ГП10 №3 с падением 75.xlsx")
    with app.app_context():
        lines = parse_price_catalog_by_brand(path, None)

    makes = {line["vehicle_make"] for line in lines}
    assert not any(m and m.startswith("-") for m in makes)
    assert "TOYOTA" in makes
    assert "NISSAN" in makes
    assert "HYUNDAI" in makes


def test_parse_price_catalog_by_brand_recognizes_alternate_title_marker_and_odd_column_order(app):
    """Регрессия по реальному файлу заказчика (testdata/Приложение ГП10 №3
    с падением 75.xlsx, листы "Dewoo"/"газ"): заголовок листа не "Марка
    (модель) технического средства X", а "Запчасти на автомобиль X" — до
    фикса такие листы вообще не попадали в brand_sheets и терялись целиком
    при импорте (0 позиций марки Daewoo/ГАЗ). Плюс у этих листов цена и
    артикул НЕ на тех позициях, что у остальных марок того же файла
    (Ед.изм. вместо артикула, артикула нет вовсе) — жёстко зашитая позиция
    колонки перепутала бы цену со служебным полем."""
    path = os.path.join(TESTDATA_DIR, "Приложение ГП10 №3 с падением 75.xlsx")
    with app.app_context():
        lines = parse_price_catalog_by_brand(path, None)

    daewoo = [l for l in lines if l["vehicle_make"] == "DAEWOO"]
    gaz = [l for l in lines if l["vehicle_make"] == "GAZ"]
    assert len(daewoo) > 0
    assert len(gaz) > 0

    # У этих двух листов в принципе нет колонки с артикулом — не должно
    # задваивать туда цену/номер по умолчанию под видом артикула.
    assert all(l["article"] is None for l in daewoo)
    assert all(l["article"] is None for l in gaz)
    # А цена — настоящая цена (число из своей колонки), не строка "шт."/"к-т".
    assert all(l["price"] is not None and l["price"] > 0 for l in daewoo)
    assert all(l["price"] is not None and l["price"] > 0 for l in gaz)


def test_sheet_title_detected_even_when_not_in_first_column(app):
    """testdata/.../Dewoo: 'Запчасти на автомобиль Daewoo Nexia' лежит во
    ВТОРОЙ ячейке первой строки (первая — пустая), не в колонке A — раньше
    заголовок листа читался только из первой ячейки."""
    path = os.path.join(TESTDATA_DIR, "Приложение ГП10 №3 с падением 75.xlsx")
    with app.app_context():
        lines = parse_price_catalog_by_brand(path, "DAEWOO")
    assert lines is not None
    assert len(lines) > 0
    assert all(l["vehicle_make"] == "DAEWOO" for l in lines if l["name"])


@pytest.mark.parametrize(
    "cyrillic,expected_latin",
    [
        ("Чери", "CHERY"),
        ("Джили", "GEELY"),
        ("Хавал", "HAVAL"),
        ("Дунфэн", "DONGFENG"),
        ("Ссангйонг", "SSANGYONG"),
        ("Пежо", "PEUGEOT"),
        ("Ленд Ровер", "LAND ROVER"),
    ],
)
def test_normalize_brand_label_covers_common_market_brands(app, cyrillic, expected_latin):
    """Заказчик не ограничится маркой из уже присланных файлов — справочник
    должен покрывать массовые марки рынка РФ (китайские бренды быстро
    растут долю продаж), не только то, что уже встретилось."""
    with app.app_context():
        assert _normalize_brand_label(cyrillic) == expected_latin


def test_parse_price_catalog_single_sheet_sections_real_customer_file(app):
    """testdata/Приложение №1 ИП Даянова З.Р..xlsx — реальный файл
    заказчика: ОДИН лист, 25000+ строк, разделы марок не отдельными
    листами (см. parse_price_catalog_by_brand), а строками-маркерами
    внутри листа (LADA (ВАЗ) / УАЗ / ГАЗ / ПАЗ / TOYOTA / GM (Шевроле,
    Опель) / Неоригинальные запчасти / Масла...), причём шапка колонок —
    третья строка листа, а не первая. Это и есть каталог, парный
    реальному тестовому заказ-наряду (волжская ркб.xlsx, CHEVROLET
    LACETTI) из этой же поставки тестовых файлов."""
    path = os.path.join(TESTDATA_DIR, "Приложение №1 ИП Даянова З.Р..xlsx")

    # До фикса это падало DocumentParseError — шапка pandas по умолчанию
    # берёт первую строку листа ("Приложение №1"), под алиасы не подходит.
    with pytest.raises(DocumentParseError):
        parse_document(path)

    with app.app_context():
        lines = parse_price_catalog_single_sheet_sections(path)

    assert lines is not None
    assert len(lines) > 20000

    by_make = {}
    for line in lines:
        by_make.setdefault(line["vehicle_make"], 0)
        by_make[line["vehicle_make"]] += 1

    # GM (Шевроле, Опель) -> CHEVROLET — ровно та марка, что в реальном
    # заказ-наряде на CHEVROLET LACETTI из этой же поставки.
    assert by_make.get("CHEVROLET", 0) > 0
    assert by_make.get("LADA", 0) > 10000  # больше половины каталога — ВАЗ
    assert by_make.get("TOYOTA", 0) > 0

    # "Неоригинальные запчасти"/"Масла..." — общие для всех марок разделы,
    # не должны превращаться в фиктивную "марку", по которой ни один
    # реальный заказ-наряд никогда не совпадёт.
    assert by_make.get(None, 0) > 0
    assert not any(m and "МАСЛ" in m for m in by_make)
    assert not any(m and "НЕОРИГИНАЛ" in m for m in by_make)

    sample = next(line for line in lines if line["vehicle_make"] == "CHEVROLET")
    assert sample["price"] is not None
    assert sample["price"] > 0


def test_parse_price_catalog_single_sheet_sections_returns_none_for_unrelated_file():
    """Обычный (не однолистовой-многобрендовый) файл не должен ошибочно
    распознаваться этим парсером — иначе он "съест" файл раньше, чем до
    него дойдёт правильный, более специфичный парсер."""
    path = os.path.join(TESTDATA_DIR, "тест 1 (исходник).xlsx")
    assert parse_price_catalog_single_sheet_sections(path) is None
